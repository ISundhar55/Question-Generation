"""
pdf_parser.py
-------------
Extracts text from PDF and DOCX files, then splits into topic-aware,
page-tagged chunks. Also extracts embedded images/diagrams/charts as
separate retrievable "image chunks" so the app can answer picture-based
questions and show teachers exactly which page an image came from.

Chunking strategy (priority order):
  1. Topic boundary detection — split at heading patterns (Chapter, Topic,
     Learning Objective, numbered sections, ALL-CAPS lines, etc.)
  2. 500-token hard cap — if a topic section is too long, slide-split it
     with 50-token overlap.
  3. Fallback — if no headings detected, pure sliding-window 500/50.

Every chunk carries: chapter, topic, text, page, chunk_type ("text" | "image"),
and (for image chunks) image_path — so every generated question can be traced
back to an exact file + page, and image-based questions can show the source
image inline.
"""

import io
import os
import re
import uuid
from typing import Optional

import docx
import fitz
from services.llm import transcribe_page_image, describe_image


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

DATA_DIR = os.path.join(os.path.dirname(__file__), "..", "data")
IMAGES_DIR = os.path.join(DATA_DIR, "images")

# Sentinel line inserted between pages so downstream chunking can recover
# which page any given line of text came from, without restructuring the
# heading-detection algorithm around a list-of-pages data model.
_PAGE_MARKER_RE = re.compile(r"^<<<PAGE:(\d+)>>>$")


def _page_marker(page_num: int) -> str:
    return f"<<<PAGE:{page_num}>>>"


# ---------------------------------------------------------------------------
# Image extraction (raw image bytes are preserved on disk for later display —
# not just described text — so a picture-based question can show its source)
# ---------------------------------------------------------------------------

# Skip tiny images (logos, bullet icons, decorative rules) — not useful for
# picture-based questions and just adds noise to the image chunk pool.
MIN_IMAGE_DIMENSION = 120


def _ensure_images_dir(doc_id: str) -> str:
    path = os.path.join(IMAGES_DIR, doc_id)
    os.makedirs(path, exist_ok=True)
    return path


def extract_images_from_pdf(file_bytes: bytes, doc_id: str) -> list[dict]:
    """
    Extract meaningful embedded images from a PDF, save them to
    data/images/{doc_id}/, and return one "image chunk" record per image:
      {chapter, topic, text, page, chunk_type: "image", image_path}

    `image_path` is a relative path (doc_id/filename.png) suitable for
    building a servable URL — never an absolute filesystem path.
    `text` is a short AI-generated caption used for embedding/search so the
    image can be retrieved when a teacher asks for a diagram/chart question.
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    out_dir = _ensure_images_dir(doc_id)
    image_chunks: list[dict] = []

    for page_index, page in enumerate(doc):
        page_num = page_index + 1
        seen_xrefs = set()

        for img_info in page.get_images(full=True):
            xref = img_info[0]
            if xref in seen_xrefs:
                continue
            seen_xrefs.add(xref)

            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                ext = base_image.get("ext", "png")
                width = base_image.get("width", 0)
                height = base_image.get("height", 0)

                if width < MIN_IMAGE_DIMENSION or height < MIN_IMAGE_DIMENSION:
                    continue  # skip icons/decorative graphics

                filename = f"page{page_num}_{uuid.uuid4().hex[:8]}.{ext}"
                out_path = os.path.join(out_dir, filename)
                with open(out_path, "wb") as f:
                    f.write(image_bytes)

                # Short caption for embedding + retrieval — not a full page
                # transcription, just enough to make the image searchable.
                try:
                    caption = describe_image(image_bytes, mime_type=f"image/{ext}")
                except Exception:
                    caption = f"Diagram/image on page {page_num}."

                image_chunks.append({
                    "chapter": "General",       # back-filled by caller using page position
                    "topic": "General",
                    "text": caption,
                    "page": page_num,
                    "chunk_type": "image",
                    "image_path": f"{doc_id}/{filename}",
                })
            except Exception as e:
                print(f"[pdf_parser] [WARNING] Failed to extract image xref={xref} on page {page_num}: {e}")

    return image_chunks


# ---------------------------------------------------------------------------
# Text extraction (page-tagged)
# ---------------------------------------------------------------------------

def _extract_page_text_with_tables(page) -> str:
    """
    Extract text from a PyMuPDF page, detecting tables and formatting them
    as Markdown tables while removing the duplicate raw text from the table cells
    to prevent double parsing.
    """
    try:
        tables = page.find_tables()
    except Exception as e:
        print(f"[pdf_parser] [WARNING] Table detection failed on page: {e}")
        return page.get_text() or ""

    if not tables or not tables.tables:
        return page.get_text() or ""

    # Get all blocks (b[4] is the text content)
    blocks = page.get_text("blocks")
    # A block is: (x0, y0, x1, y1, "text", block_no, block_type)

    table_rects = [fitz.Rect(t.bbox) for t in tables.tables]

    # Filter out blocks that are substantially inside any table rect
    non_table_blocks = []
    for b in blocks:
        block_rect = fitz.Rect(b[0], b[1], b[2], b[3])
        overlaps_table = False
        for trect in table_rects:
            intersection = block_rect & trect
            if not intersection.is_empty:
                intersect_area = intersection.get_area()
                block_area = block_rect.get_area()
                # If block overlaps by more than 50%, skip it as duplicate table cell text
                if block_area > 0 and (intersect_area / block_area) > 0.5:
                    overlaps_table = True
                    break
        if not overlaps_table:
            non_table_blocks.append(b)

    # Combine non-table text blocks and Markdown table representations
    # Sort them vertically by coordinate to maintain readable page layout
    items = []
    for b in non_table_blocks:
        items.append((b[1], b[0], b[4].strip(), "block"))

    for t in tables.tables:
        try:
            md_table = t.to_markdown()
            if md_table and md_table.strip():
                items.append((t.bbox[1], t.bbox[0], md_table.strip(), "table"))
        except Exception as e:
            print(f"[pdf_parser] [WARNING] Failed to convert table to markdown: {e}")

    # Sort by y0 (top coordinate) then x0 (left coordinate)
    items.sort(key=lambda x: (x[0], x[1]))

    return "\n\n".join(item[2] for item in items if item[2])


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file.
    Uses fast local extraction for normal text pages; falls back to Gemini
    multimodal transcription only for pages with little/no extractable text
    but visible images/tables (scanned pages, rich diagrams).

    Each page's text is wrapped with a <<<PAGE:N>>> marker line so that
    chunk_text() can tag every resulting chunk with its source page.
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []

    print(f"[pdf_parser] Processing PDF with {len(doc)} pages")

    for i, page in enumerate(doc):
        page_num = i + 1

        local_text = (page.get_text() or "").strip()

        has_images = len(page.get_images()) > 0
        has_tables = False
        try:
            has_tables = len(page.find_tables().tables) > 0
        except Exception:
            pass

        # Only fall back to Gemini multimodal OCR if the page has little
        # extractable text (likely scanned) but visible images/tables.
        is_scanned_ocr = len(local_text) < 150 and (has_images or has_tables)

        if not is_scanned_ocr:
            print(f"[pdf_parser] Page {page_num}/{len(doc)} extracted locally in milliseconds.")
            page_text = _extract_page_text_with_tables(page)
        else:
            print(f"[pdf_parser] Page {page_num}/{len(doc)} has low extractable text but contains images/tables. Ingesting via Gemini Multimodal...")
            try:
                pix = page.get_pixmap(dpi=150)
                image_bytes = pix.tobytes("png")
                page_text = transcribe_page_image(image_bytes)
            except Exception as e:
                print(f"[pdf_parser] [WARNING] Page {page_num} multimodal failed: {e}. Falling back to local text extraction.")
                page_text = local_text

        pages.append(f"{_page_marker(page_num)}\n{page_text}")

    return "\n\n".join(pages)


def _format_docx_table_as_markdown(table) -> str:
    """Format a python-docx Table object as a Markdown table."""
    rows = []
    for row in table.rows:
        row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(row_text)

    if not rows:
        return ""

    col_count = len(rows[0])
    markdown_lines = []
    markdown_lines.append("| " + " | ".join(rows[0]) + " |")
    markdown_lines.append("| " + " | ".join(["---"] * col_count) + " |")
    for row in rows[1:]:
        if len(row) < col_count:
            row.extend([""] * (col_count - len(row)))
        elif len(row) > col_count:
            row = row[:col_count]
        markdown_lines.append("| " + " | ".join(row) + " |")

    return "\n" + "\n".join(markdown_lines) + "\n"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract all text from a DOCX file, including tables in Markdown format.
    DOCX has no fixed "pages" (pagination is a rendering concern, not a
    stored property), so every chunk is tagged page=None for this format —
    the UI shows the filename/chapter for citation instead of a page number.
    """
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    doc = docx.Document(io.BytesIO(file_bytes))
    body_elements = []

    for element in doc.element.body:
        if element.tag.endswith('p'):
            p = Paragraph(element, doc)
            if p.text.strip():
                body_elements.append(p.text.strip())
        elif element.tag.endswith('tbl'):
            table = Table(element, doc)
            markdown_table = _format_docx_table_as_markdown(table)
            if markdown_table:
                body_elements.append(markdown_table)

    return "\n\n".join(body_elements)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    elif ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file format: .{ext}")


# ---------------------------------------------------------------------------
# Heading detection
# ---------------------------------------------------------------------------
#
# NOTE: patterns accept an optional hyphen/dash/em-dash between "chapter"/
# "unit"/"lesson" and the number (e.g. "CHAPTER-4", not just "Chapter 4").
# Real-world syllabus PDFs commonly use the hyphenated style, and the
# original space-only pattern silently failed to detect chapter boundaries
# for those files — every chunk fell back to the "Introduction" default,
# which broke chapter-name display in citations (functionally chapters
# still filtered correctly via the keyword fallback in metadata_store, but
# the displayed label was wrong). Verified against a real 36-page syllabus.

_HEADING_PATTERNS = [
    re.compile(r"^\s*chapter[\s\-–—]*\d+", re.IGNORECASE),
    re.compile(r"^\s*unit[\s\-–—]*\d+", re.IGNORECASE),
    re.compile(r"^\s*topic\s*[:\-–]", re.IGNORECASE),
    re.compile(r"^\s*lesson[\s\-–—]*\d+", re.IGNORECASE),
    re.compile(r"^\s*learning\s+objective", re.IGNORECASE),
    re.compile(r"^\s*\d+\.\s+[A-Z]"),           # "1. Introduction"
    re.compile(r"^\s*\d+\.\d+\s+[A-Za-z]"),     # "1.1 Fractions"
    re.compile(r"^[A-Z][A-Z\s'’\-]{5,}$"),       # ALL CAPS heading (apostrophes/hyphens OK)
]

_CHAPTER_PATTERNS = [
    re.compile(r"^\s*chapter[\s\-–—]*\d+", re.IGNORECASE),
    re.compile(r"^\s*unit[\s\-–—]*\d+", re.IGNORECASE),
]

# Words that commonly appear at the end of a truncated heading line when the
# PDF text extractor splits a multi-word title across physical lines.
# Detecting these lets chunk_text() merge the next short line before locking
# in the chapter/topic label — e.g. "5. Integration of" + "Knowledge and Ideas"
# becomes "5. Integration of Knowledge and Ideas" instead of "5. Integration of".
_INCOMPLETE_HEADING_WORDS = {
    "of", "and", "the", "in", "to", "a", "an", "for", "with", "at", "by",
    "or", "but", "from", "on", "as", "into", "over", "about", "&",
    "its", "their", "that", "which", "between", "among",
}


def _is_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    return any(p.match(stripped) for p in _HEADING_PATTERNS)


def _is_chapter(line: str) -> bool:
    stripped = line.strip()
    return any(p.match(stripped) for p in _CHAPTER_PATTERNS)


def _heading_is_incomplete(heading: str) -> bool:
    """Return True if a heading ends with a dangling word or a colon, indicating
    the PDF split a multi-word title across physical lines.

    Examples that return True:
      '5. Integration of'             ← ends with preposition
      'Chapter 6: Integration of'     ← ends with preposition
      'knowledge and ideas:'          ← trailing colon = explicit continuation
    """
    stripped = heading.strip()
    if not stripped:
        return False
    # A trailing colon explicitly signals that more text follows on the next line
    # (e.g. "Integration of knowledge and ideas:" → "Long passage practice")
    if stripped.endswith(':'):
        return True
    words = stripped.split()
    last = words[-1].lower().rstrip(".,;:\"'")
    return last in _INCOMPLETE_HEADING_WORDS


# ---------------------------------------------------------------------------
# Tokenisation (word-level, simple)
# ---------------------------------------------------------------------------

def _word_count(text: str) -> int:
    return len(text.split())


# ---------------------------------------------------------------------------
# Sliding-window split for oversized sections
# ---------------------------------------------------------------------------

MAX_TOKENS = 500
OVERLAP_TOKENS = 50


def _split_long_text(text: str, chapter: str, topic: str, page: Optional[int]) -> list[dict]:
    """Split text that exceeds MAX_TOKENS with overlap. Page is inherited
    (approximate — the section's first page) since precise per-word page
    boundaries aren't tracked at this granularity."""
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + MAX_TOKENS, len(words))
        chunk_text = " ".join(words[start:end])
        chunks.append({
            "chapter": chapter, "topic": topic, "text": chunk_text,
            "page": page, "chunk_type": "text",
        })
        if end == len(words):
            break
        start = end - OVERLAP_TOKENS
    return chunks


# ---------------------------------------------------------------------------
# Topic-aware chunking
# ---------------------------------------------------------------------------

def chunk_text(text: str) -> list[dict]:
    """
    Split text into chunks using heading detection first, then enforce the
    500-token max with 50-token overlap. Tracks <<<PAGE:N>>> markers (if
    present — PDFs only) so every chunk carries its source page number.

    A bare chapter-number heading ("CHAPTER-4") immediately followed by its
    descriptive title line ("CELLS AND ORGANISMS") is merged into one
    readable chapter label ("CHAPTER-4 — CELLS AND ORGANISMS") rather than
    showing the bare number in citations.

    Returns list of dicts: {chapter, topic, text, page, chunk_type: "text"}
    """
    lines = text.splitlines()

    sections = []          # [{chapter, topic, lines, page}]
    current_chapter = "Introduction"
    current_topic = "General"
    current_lines: list[str] = []
    current_page: Optional[int] = None
    section_start_page: Optional[int] = None
    pending_chapter_number: Optional[str] = None
    # True when the last detected topic heading ended with a dangling word
    # (e.g. "of", "and") and the next short non-heading line should be merged
    # into the topic label rather than treated as body text.
    pending_topic_continuation: bool = False

    for line in lines:
        marker = _PAGE_MARKER_RE.match(line.strip())
        if marker:
            current_page = int(marker.group(1))
            continue

        if _is_heading(line):
            stripped = line.strip()
            pending_topic_continuation = False  # a new heading resets any pending merge

            if _is_chapter(stripped):
                if current_lines:
                    sections.append({
                        "chapter": current_chapter, "topic": current_topic,
                        "lines": current_lines, "page": section_start_page,
                    })
                    current_lines = []
                pending_chapter_number = stripped
                current_chapter = stripped   # placeholder until/unless a title line follows
                current_topic = stripped
                section_start_page = current_page
                # A chapter heading can itself be split across lines
                # (e.g. "Chapter 6: Integration of" + "knowledge and ideas:").
                # Apply the same incomplete-heading merging that topic headings use.
                pending_topic_continuation = _heading_is_incomplete(current_chapter)
            else:
                if pending_chapter_number and not current_lines:
                    # This heading immediately follows a bare chapter number
                    # with no body text yet — treat it as that chapter's title.
                    current_chapter = f"{pending_chapter_number} — {stripped}"
                    current_topic = stripped
                    pending_chapter_number = None
                else:
                    if current_lines:
                        sections.append({
                            "chapter": current_chapter, "topic": current_topic,
                            "lines": current_lines, "page": section_start_page,
                        })
                        current_lines = []
                    current_topic = stripped
                    pending_chapter_number = None
                section_start_page = current_page
                # Flag if this topic heading ends mid-phrase so the next line can extend it
                pending_topic_continuation = _heading_is_incomplete(current_topic)
        else:
            if line.strip():
                stripped_body = line.strip()
                # If the previous heading ended mid-phrase (e.g. "5. Integration of")
                # and no body lines have been collected yet, treat this short line as
                # the continuation of the heading rather than as body text.
                if pending_topic_continuation and not current_lines and len(stripped_body) <= 100:
                    old_topic = current_topic
                    current_topic = f"{current_topic} {stripped_body}"
                    # Extend current_chapter if it was formed from or equals the old topic label.
                    if current_chapter.endswith(old_topic):
                        current_chapter = current_chapter[:-len(old_topic)] + current_topic
                    # Also keep pending_chapter_number in sync so the chapter-title
                    # merge logic (pending_chapter_number + next heading) stays correct.
                    if pending_chapter_number == old_topic:
                        pending_chapter_number = current_topic
                    # Keep merging if the extended label still ends mid-phrase.
                    pending_topic_continuation = _heading_is_incomplete(current_topic)
                else:
                    pending_topic_continuation = False
                    if not current_lines:
                        section_start_page = current_page
                    current_lines.append(line)
                    pending_chapter_number = None

    if current_lines:
        sections.append({
            "chapter": current_chapter, "topic": current_topic,
            "lines": current_lines, "page": section_start_page,
        })

    # Enforce 500-token max
    final_chunks: list[dict] = []
    for section in sections:
        section_text = " ".join(section["lines"])
        if _word_count(section_text) <= MAX_TOKENS:
            if section_text.strip():
                final_chunks.append({
                    "chapter": section["chapter"],
                    "topic": section["topic"],
                    "text": section_text.strip(),
                    "page": section["page"],
                    "chunk_type": "text",
                })
        else:
            final_chunks.extend(
                _split_long_text(section_text, section["chapter"], section["topic"], section["page"])
            )

    # Fallback: if no structure found at all, sliding window
    if not final_chunks:
        final_chunks = _split_long_text(text, "General", "General", None)

    return final_chunks
